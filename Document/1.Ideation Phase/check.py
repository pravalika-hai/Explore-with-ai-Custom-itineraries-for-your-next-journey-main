import os
import glob
from docx import Document
import re

directory = r"c:\Users\sriha\Downloads\Explore-with-ai-Custom-itineraries-for-your-next-journey-main\Explore-with-ai-Custom-itineraries-for-your-next-journey-main\Document"
docx_files = glob.glob(os.path.join(directory, "**", "*.docx"), recursive=True)

pattern = re.compile(r"LTVIP\w+")

for filepath in docx_files:
    try:
        doc = Document(filepath)
        found = set()
        for p in doc.paragraphs:
            matches = pattern.findall(p.text)
            for m in matches:
                found.add(m)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        matches = pattern.findall(p.text)
                        for m in matches:
                            found.add(m)
        if found:
            print(f"{filepath}: {found}")
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
