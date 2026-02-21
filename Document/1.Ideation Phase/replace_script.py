import os
import glob
from docx import Document

def replace_text_in_docx(filepath, search_text, replace_text):
    try:
        doc = Document(filepath)
        changed = False

        # Replace in paragraphs
        for p in doc.paragraphs:
            if search_text in p.text:
                for run in p.runs:
                    if search_text in run.text:
                        run.text = run.text.replace(search_text, replace_text)
                        changed = True
                # If text was split across runs, a simple replace on runs might fail.
                # Here's a naive approach for block-level replace which might lose some formatting,
                # but let's try run-level first and then block-level fallback if needed.
                if search_text in p.text:
                    p.text = p.text.replace(search_text, replace_text)
                    changed = True

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if search_text in p.text:
                            for run in p.runs:
                                if search_text in run.text:
                                    run.text = run.text.replace(search_text, replace_text)
                                    changed = True
                            if search_text in p.text:
                                p.text = p.text.replace(search_text, replace_text)
                                changed = True
        
        # Replace in sections (Headers/Footers)
        for section in doc.sections:
            for p in section.header.paragraphs:
                if search_text in p.text:
                    for run in p.runs:
                        if search_text in run.text:
                            run.text = run.text.replace(search_text, replace_text)
                            changed = True
                    if search_text in p.text:
                        p.text = p.text.replace(search_text, replace_text)
                        changed = True

            for p in section.footer.paragraphs:
                if search_text in p.text:
                    for run in p.runs:
                        if search_text in run.text:
                            run.text = run.text.replace(search_text, replace_text)
                            changed = True
                    if search_text in p.text:
                        p.text = p.text.replace(search_text, replace_text)
                        changed = True


        if changed:
            doc.save(filepath)
            print(f"Updated: {filepath}")
        else:
            print(f"No changes needed: {filepath}")

    except Exception as e:
        print(f"Error processing {filepath}: {e}")

directory = r"c:\Users\sriha\Downloads\Explore-with-ai-Custom-itineraries-for-your-next-journey-main\Explore-with-ai-Custom-itineraries-for-your-next-journey-main\Document"
docx_files = glob.glob(os.path.join(directory, "**", "*.docx"), recursive=True)

search_str = "LTVIP2026TMIDS85836"
replace_str = "LTVIP2026TMIDS88533"

for f in docx_files:
    replace_text_in_docx(f, search_str, replace_str)
